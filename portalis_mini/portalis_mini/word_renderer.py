from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Dict, Optional

from docx import Document

from portalis_models import (
    CrewRecord,
    GenerationResult,
    RendererDefinition,
    VesselRecord,
    VoyageContext,
)


class WordRendererError(Exception):
    """Raised when a Word renderer operation fails."""


class WordRenderer:
    def __init__(self, templates_root: str | Path, output_root: str | Path) -> None:
        self.templates_root = Path(templates_root)
        self.output_root = Path(output_root)
        self.output_root.mkdir(parents=True, exist_ok=True)

    def render_crew_landing_permit(
        self,
        renderer: RendererDefinition,
        crew: CrewRecord,
        vessel: VesselRecord,
        voyage: VoyageContext,
    ) -> GenerationResult:
        template_path = self._resolve_template_path(renderer.template_path)
        output_path = self._build_output_path(
            crew_id=crew.crew_id,
            family_name=crew.family_name,
            given_name=crew.given_name,
            suffix="_CLP.docx",
        )

        field_map = self._build_clp_field_map(crew=crew, vessel=vessel, voyage=voyage)

        doc = Document(template_path)
        self._replace_bookmarks_like_placeholders(doc, field_map)
        doc.save(output_path)

        return GenerationResult(
            output_id=f"{crew.crew_id}_clp_{datetime.utcnow().strftime('%Y%m%d%H%M%S')}",
            renderer_id=renderer.renderer_id,
            output_path=str(output_path),
            generated_at=datetime.utcnow(),
            source_entity_ids=[crew.crew_id, vessel.vessel_id, voyage.voyage_id],
            warnings=[],
            review_items=[],
        )

    def _resolve_template_path(self, template_path: str) -> Path:
        candidate = Path(template_path)
        if candidate.is_absolute():
            resolved = candidate
        else:
            resolved = self.templates_root / template_path

        if not resolved.exists():
            raise WordRendererError(f"Template not found: {resolved}")

        return resolved

    def _build_output_path(
        self,
        crew_id: str,
        family_name: str,
        given_name: str,
        suffix: str,
    ) -> Path:
        safe_name = self._safe_filename(f"{crew_id}_{family_name}_{given_name}{suffix}")
        return self.output_root / safe_name

    def _safe_filename(self, value: str) -> str:
        bad_chars = ['\\', '/', ':', '*', '?', '\"', '<', '>', '|']
        for ch in bad_chars:
            value = value.replace(ch, "_")
        return value

    def _build_clp_field_map(
        self,
        crew: CrewRecord,
        vessel: VesselRecord,
        voyage: VoyageContext,
    ) -> Dict[str, str]:
        passport = crew.passports[0] if crew.passports else None

        return {
            "FINNo": "",
            "FamilyName": crew.family_name or "",
            "GivenName": crew.given_name or "",
            "Initial": crew.middle_initial or "",
            "HomeAddress": crew.notes or "",
            "AddressUS": f"C/O M/V {vessel.name}, Port of {voyage.current_port}" if voyage.current_port else f"C/O M/V {vessel.name}",
            "Hair": "",
            "Eyes": "",
            "Height": "",
            "Weight": "",
            "DateOfBirth": crew.date_of_birth.isoformat() if crew.date_of_birth else "",
            "PlaceOfBirth": crew.place_of_birth or "",
            "PassportNumberNationality": self._passport_nat_line(passport, crew.nationality),
        }

    def _passport_nat_line(self, passport, nationality: Optional[str]) -> str:
        if passport and passport.number and nationality:
            return f"{passport.number} / {nationality}"
        if passport and passport.number:
            return passport.number
        return nationality or ""

    def _replace_bookmarks_like_placeholders(self, doc: Document, field_map: Dict[str, str]) -> None:
        # First version: placeholder text replacement like [FamilyName]
        for paragraph in doc.paragraphs:
            for key, value in field_map.items():
                token = f"[{key}]"
                if token in paragraph.text:
                    for run in paragraph.runs:
                        if token in run.text:
                            run.text = run.text.replace(token, value)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in field_map.items():
                            token = f"[{key}]"
                            if token in paragraph.text:
                                for run in paragraph.runs:
                                    if token in run.text:
                                        run.text = run.text.replace(token, value)